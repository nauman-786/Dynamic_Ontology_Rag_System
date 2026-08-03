import os
import fitz  # PyMuPDF
from docx import Document
from documents.metadata import ParsedDocument, DocumentMetadata

class DocumentLoader:
    """Handles parsing of various file types into a unified ParsedDocument format."""
    
    @staticmethod
    def load_document(file_path: str) -> ParsedDocument:
        """Determines file type and routes to the correct parser."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
            
        ext = os.path.splitext(file_path)[1].lower()
        filename = os.path.basename(file_path)
        
        if ext == ".pdf":
            return DocumentLoader._load_pdf(file_path, filename)
        elif ext == ".docx":
            return DocumentLoader._load_docx(file_path, filename)
        elif ext == ".txt":
            return DocumentLoader._load_txt(file_path, filename)
        else:
            raise ValueError(f"Unsupported file extension: {ext}")

    @staticmethod
    def _load_pdf(file_path: str, filename: str) -> ParsedDocument:
        text = ""
        total_pages = 0
        with fitz.open(file_path) as pdf_doc:
            total_pages = len(pdf_doc)
            for page in pdf_doc:
                text += page.get_text("text") + "\n\n"
                
        metadata = DocumentMetadata(
            filename=filename,
            extension=".pdf",
            total_pages=total_pages
        )
        return ParsedDocument(text=text.strip(), metadata=metadata)

    @staticmethod
    def _load_docx(file_path: str, filename: str) -> ParsedDocument:
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        
        metadata = DocumentMetadata(
            filename=filename,
            extension=".docx",
            total_pages=1 # DOCX page count requires rendering, defaulting to 1
        )
        return ParsedDocument(text=text.strip(), metadata=metadata)

    @staticmethod
    def _load_txt(file_path: str, filename: str) -> ParsedDocument:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
            
        metadata = DocumentMetadata(
            filename=filename,
            extension=".txt",
            total_pages=1
        )
        return ParsedDocument(text=text.strip(), metadata=metadata)